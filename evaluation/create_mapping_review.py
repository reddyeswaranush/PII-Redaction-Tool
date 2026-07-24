"""Create a human-review CSV from the detector's persisted mapping.

The CSV deliberately contains original PII values and must remain local. Review
each row by setting ``is_actual_pii`` to ``yes`` or ``no`` and, when relevant,
fill ``entity_type`` and ``notes``.
"""

from __future__ import annotations

import argparse
import csv
import json
from pathlib import Path

from docx import Document

from src.config import INPUT_FILE, MAPPING_FILE, OUTPUT_FILE


def _document_text(path: Path) -> str:
    """Extract searchable paragraph text from a DOCX document."""
    document = Document(str(path))
    chunks: list[str] = []
    for paragraph in document.paragraphs:
        chunks.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(chunks)


def _context(text: str, value: str, radius: int = 100) -> str:
    """Return the first source context containing a mapped value."""
    position = text.casefold().find(value.casefold())
    if position < 0:
        return ""
    start = max(0, position - radius)
    end = min(len(text), position + len(value) + radius)
    return text[start:end].replace("\r", " ").replace("\n", " ")


def create_review(
    mapping_path: Path,
    input_path: Path,
    output_path: Path,
    review_path: Path,
) -> None:
    """Write one review row for each unique mapping entry."""
    mapping = json.loads(mapping_path.read_text(encoding="utf-8"))
    original_text = _document_text(input_path)
    redacted_text = _document_text(output_path)
    rows: list[dict[str, object]] = []

    for original, replacement in sorted(mapping.items(), key=lambda item: item[0].casefold()):
        rows.append({
            "original": original,
            "replacement": replacement,
            "original_occurrences": original_text.casefold().count(original.casefold()),
            "replacement_occurrences": redacted_text.casefold().count(replacement.casefold()),
            "source_context": _context(original_text, original),
            "is_actual_pii": "",
            "entity_type": "",
            "notes": "",
        })

    review_path.parent.mkdir(parents=True, exist_ok=True)
    with review_path.open("w", encoding="utf-8", newline="") as file:
        writer = csv.DictWriter(file, fieldnames=list(rows[0]))
        writer.writeheader()
        writer.writerows(rows)
    print(f"Wrote {len(rows)} review rows to {review_path}")


def main() -> None:
    """Create the local mapping review file."""
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--mapping", type=Path, default=MAPPING_FILE)
    parser.add_argument("--input", type=Path, default=INPUT_FILE)
    parser.add_argument("--output", type=Path, default=OUTPUT_FILE)
    parser.add_argument("--review", type=Path, default=Path("evaluation/mapping_review.csv"))
    args = parser.parse_args()
    create_review(args.mapping, args.input, args.output, args.review)


if __name__ == "__main__":
    main()
