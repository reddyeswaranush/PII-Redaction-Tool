"""DOCX traversal, persistence, statistics, and formatting-safe edits."""

from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Dict, Iterable, Iterator, List, TYPE_CHECKING

from docx import Document
from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph
from docx.text.run import Run

try:
    from docx.table import Table, _Cell as Cell
except ImportError:  # pragma: no cover - compatibility with older python-docx
    from docx.table import Cell, Table

from .config import MAPPING_FILE, OUTPUT_DIR

if TYPE_CHECKING:
    from .detector import PIIEntity

LOGGER = logging.getLogger(__name__)


def create_output_directory() -> None:
    """Create the configured output directory."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def load_document(file_path: Path) -> DocumentType:
    """Load a DOCX file."""
    return Document(str(file_path))


def save_document(document: DocumentType, output_path: Path) -> None:
    """Save a DOCX file, creating its parent directory when necessary."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def get_paragraphs(document: DocumentType) -> List[Paragraph]:
    """Return top-level body paragraphs (legacy API)."""
    return list(document.paragraphs)


def get_tables(document: DocumentType) -> List[Table]:
    """Return top-level body tables (legacy API)."""
    return list(document.tables)


def iter_table_cells(
    tables: Iterable[Table], seen: set[int] | None = None,
) -> Iterator[Cell]:
    """Yield each table cell once, including arbitrarily nested tables."""
    seen = seen if seen is not None else set()
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                marker = id(cell._tc)
                if marker in seen:
                    continue
                seen.add(marker)
                yield cell
                yield from iter_table_cells(cell.tables, seen)


def get_table_cells(document: DocumentType) -> Iterator[Cell]:
    """Yield unique cells in all top-level document tables."""
    yield from iter_table_cells(document.tables)


def iter_runs(paragraph: Paragraph) -> Iterator[Run]:
    """Yield normal and hyperlink runs without changing their XML wrappers."""
    inner = getattr(paragraph, "iter_inner_content", None)
    if inner is None:
        yield from paragraph.runs
        return
    yielded = False
    for item in inner():
        if isinstance(item, Run):
            yielded = True
            yield item
        elif hasattr(item, "runs"):
            for run in item.runs:
                yielded = True
                yield run
    if not yielded:
        yield from paragraph.runs


def iter_cell_paragraphs(cell: Cell) -> Iterator[Paragraph]:
    """Yield paragraphs directly in a cell and recursively nested tables."""
    yield from cell.paragraphs
    for nested_cell in iter_table_cells(cell.tables):
        yield from nested_cell.paragraphs


def _story_parts(document: DocumentType) -> Iterator[object]:
    """Yield document and all header/footer stories, including variants."""
    yield document
    seen: set[int] = set()
    for section in document.sections:
        for story in (
            section.header, section.first_page_header, section.even_page_header,
            section.footer, section.first_page_footer, section.even_page_footer,
        ):
            marker = id(story._element)
            if marker not in seen:
                seen.add(marker)
                yield story


def iter_document_paragraphs(document: DocumentType) -> Iterator[Paragraph]:
    """Yield unique paragraphs in body, tables, headers, and footers."""
    seen_paragraphs: set[int] = set()
    seen_cells: set[int] = set()
    for story in _story_parts(document):
        paragraphs = list(story.paragraphs)
        cells = list(iter_table_cells(story.tables, seen_cells))
        for cell in cells:
            paragraphs.extend(iter_cell_paragraphs(cell))
        for paragraph in paragraphs:
            marker = id(paragraph._p)
            if marker not in seen_paragraphs:
                seen_paragraphs.add(marker)
                yield paragraph


def replace_entities_in_paragraph(
    paragraph: Paragraph,
    entities: Iterable["PIIEntity"],
    anonymizer: object,
) -> int:
    """Replace entity spans across runs while retaining run formatting.

    A replacement is inserted into the first run containing an entity's start;
    any remaining source characters in later runs are removed. The run XML is
    retained, so bold, italic, color, font, and hyperlink wrappers survive.
    """
    runs = list(iter_runs(paragraph))
    ordered = sorted(entities, key=lambda item: (item.start, item.end))
    if not runs or not ordered:
        return 0
    offsets: list[tuple[Run, str, int, int]] = []
    offset = 0
    for run in runs:
        text = run.text or ""
        offsets.append((run, text, offset, offset + len(text)))
        offset += len(text)

    changed = 0
    for run, text, start, end in offsets:
        if not text:
            continue
        output: list[str] = []
        cursor = 0
        for entity in ordered:
            if entity.end <= start:
                continue
            if entity.start >= end:
                break
            local_start = max(entity.start, start) - start
            local_end = min(entity.end, end) - start
            if local_end <= local_start:
                continue
            if local_start > cursor:
                output.append(text[cursor:local_start])
            if start <= entity.start < end:
                output.append(anonymizer.anonymize(entity.entity_type, entity.text))
                changed += 1
            cursor = max(cursor, local_end)
        output.append(text[cursor:])
        new_text = "".join(output)
        if new_text != text:
            run.text = new_text
    return changed


def load_mapping() -> Dict[str, str]:
    """Load the persisted mapping, or return an empty mapping."""
    if not MAPPING_FILE.exists():
        return {}
    try:
        with MAPPING_FILE.open("r", encoding="utf-8") as file:
            data = json.load(file)
    except (OSError, json.JSONDecodeError) as error:
        LOGGER.warning("Could not load mapping file %s: %s", MAPPING_FILE, error)
        return {}
    return data if isinstance(data, dict) else {}


def save_mapping(mapping: Dict[str, str]) -> None:
    """Persist an original-to-fake mapping as UTF-8 JSON."""
    MAPPING_FILE.parent.mkdir(parents=True, exist_ok=True)
    with MAPPING_FILE.open("w", encoding="utf-8") as file:
        json.dump(mapping, file, indent=2, ensure_ascii=False, sort_keys=True)


def replace_text(text: str, replacements: Dict[str, str]) -> str:
    """Replace longer strings first (legacy plain-text API)."""
    for original, fake in sorted(replacements.items(), key=lambda item: -len(item[0])):
        text = text.replace(original, fake)
    return text


def count_paragraphs(document: DocumentType) -> int:
    """Count unique paragraphs across all DOCX stories."""
    return sum(1 for _ in iter_document_paragraphs(document))


def count_tables(document: DocumentType) -> int:
    """Count unique tables in body, nested cells, and header/footer stories."""
    seen: set[int] = set()
    count = 0
    for story in _story_parts(document):
        for table in story.tables:
            for cell in iter_table_cells([table]):
                for nested in [table, *cell.tables]:
                    marker = id(nested._tbl)
                    if marker not in seen:
                        seen.add(marker)
                        count += 1
    return count


def count_words(document: DocumentType) -> int:
    """Count words in unique paragraphs across all document stories."""
    return sum(len(paragraph.text.split()) for paragraph in iter_document_paragraphs(document))


def document_statistics(document: DocumentType) -> Dict[str, int]:
    """Return paragraph, table, word, and character counts."""
    paragraphs = list(iter_document_paragraphs(document))
    return {
        "paragraphs": len(paragraphs),
        "tables": count_tables(document),
        "words": sum(len(p.text.split()) for p in paragraphs),
        "characters": sum(len(p.text) for p in paragraphs),
    }


def print_document_summary(document: DocumentType) -> None:
    """Log document statistics for callers using the legacy helper."""
    for key, value in document_statistics(document).items():
        LOGGER.info("%-12s: %s", key.title(), value)
